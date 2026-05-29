"""
models/exporter.py
──────────────────
Generate downloadable DOCX and PDF files from plain-text resume content.
No Streamlit imports here — pure file-generation logic.
"""

from io import BytesIO
import textwrap

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt
from reportlab.lib import colors
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas

from utils import try_parse_json
from io import BytesIO
from reportlab.lib.pagesizes import letter
from reportlab.lib import colors
from reportlab.lib.styles import ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import (
    SimpleDocTemplate, Paragraph, Spacer, HRFlowable, KeepTogether
)
from reportlab.lib.enums import TA_LEFT, TA_CENTER


# ── Shared helpers ─────────────────────────────────────────────────────────────

def _clean_text(text: str) -> str:
    """Replace fancy unicode punctuation with ASCII equivalents."""
    replacements = {"■": "-", "•": "-", "–": "-", "—": "-"}
    for old, new in replacements.items():
        text = text.replace(old, new)
    return text


def _is_heading(line: str) -> bool:
    """Identify section headings and titles."""
    clean = line.strip()
    if not clean:
        return False
    if clean.upper() == clean and len(clean.split()) <= 8:
        return True
    if any(clean.lower().startswith(prefix) for prefix in [
        "summary", "experience", "education", "projects", "skills", "certifications",
        "contact", "professional summary", "objective", "achievements"
    ]):
        return True
    if clean.endswith(":"):
        return True
    return False


def _parse_resume_header(header_lines) -> dict:
    name = ""
    title = ""
    contact = []
    extras = []

    for line in header_lines:
        lower = line.lower()
        if any(token in lower for token in ["@", "linkedin", "github", "http", "+", "phone", "tel", "www."]):
            contact.append(line)
            continue

        parts = line.split()
        if not name and 1 <= len(parts) <= 5 and all(part.isalpha() or part.replace("'", "").isalpha() for part in parts):
            name = line
            continue

        if not title:
            title = line
            continue

        extras.append(line)

    return {
        "name": name,
        "title": title,
        "contact": contact,
        "extras": extras,
    }


def _split_header_and_sections(resume_text: str):
    lines = [line.rstrip() for line in resume_text.splitlines()]
    header_lines = []
    sections = []
    current = {"title": "", "items": []}
    found_section = False

    for line in lines:
        stripped = line.strip()
        if not stripped and not found_section:
            continue

        if not found_section and stripped and _is_heading(stripped):
            found_section = True

        if not found_section:
            if stripped:
                header_lines.append(stripped)
            continue

        if stripped and _is_heading(stripped):
            if current["title"] or current["items"]:
                sections.append(current)
            current = {"title": stripped, "items": []}
        elif stripped:
            current["items"].append(stripped)
        elif current["items"]:
            current["items"].append("")

    if current["title"] or current["items"]:
        sections.append(current)

    return header_lines, sections


def _load_structured_resume(resume_text: str):
    parsed = try_parse_json(resume_text)
    if not parsed or not isinstance(parsed, dict):
        return None
    if "final_resume" in parsed and isinstance(parsed["final_resume"], dict):
        return parsed["final_resume"]
    if "final_ats_optimized_resume" in parsed and isinstance(parsed["final_ats_optimized_resume"], dict):
        return parsed["final_ats_optimized_resume"]
    return parsed


def _flatten_value(value):
    if isinstance(value, str):
        return value
    if isinstance(value, dict):
        parts = []
        for key, item in value.items():
            if isinstance(item, (str, int, float)):
                parts.append(f"{key.replace('_', ' ').title()}: {item}")
            elif isinstance(item, list):
                parts.extend(_flatten_value(item).split("\n"))
            elif isinstance(item, dict):
                parts.extend(_flatten_value(item).split("\n"))
        return " | ".join(parts)
    if isinstance(value, list):
        parts = []
        for item in value:
            if isinstance(item, (str, int, float)):
                parts.append(str(item))
            elif isinstance(item, dict):
                nested = _flatten_value(item)
                if nested:
                    parts.append(nested)
        return " | ".join(parts)
    return str(value)


def _resume_dict_to_text(resume: dict) -> str:
    lines = []
    name = resume.get("name")
    title = resume.get("title")
    contact = resume.get("contact") or resume.get("contact_info")

    if name:
        lines.append(name)
    if title:
        lines.append(title)
    if contact:
        lines.append(_flatten_value(contact))

    lines.append("")

    sections = [
        ("OBJECTIVE", resume.get("objective")),
        ("EDUCATION", resume.get("education")),
        ("TECHNICAL SKILLS", resume.get("technical_skills")),
        ("TOOLS & PLATFORMS", resume.get("tools_and_platforms")),
        ("PROJECTS", resume.get("projects")),
        ("EXPERIENCE", resume.get("experience")),
        ("CERTIFICATIONS", resume.get("certifications")),
        ("ACHIEVEMENTS", resume.get("achievements")),
    ]

    for title, value in sections:
        if not value:
            continue
        lines.append(title)
        if isinstance(value, str):
            lines.append(value)
        elif isinstance(value, dict):
            for key, item in value.items():
                if isinstance(item, list):
                    lines.append(f"{key.replace('_', ' ').title()}:" )
                    for entry in item:
                        if isinstance(entry, dict):
                            nested = _flatten_value(entry)
                            for line in nested.split("\n"):
                                lines.append(f"- {line}")
                        else:
                            lines.append(f"- {entry}")
                elif isinstance(item, dict):
                    lines.append(f"{key.replace('_', ' ').title()}:" )
                    nested = _flatten_value(item)
                    for line in nested.split("\n"):
                        lines.append(f"- {line}")
                else:
                    lines.append(f"{key.replace('_', ' ').title()}: {item}")
        elif isinstance(value, list):
            for item in value:
                if isinstance(item, dict):
                    title_line = item.get("title") or item.get("role") or item.get("name")
                    if title_line:
                        lines.append(title_line)
                    for key, detail in item.items():
                        if key in ("title", "role", "name"):
                            continue
                        if isinstance(detail, list):
                            lines.append(f"{key.replace('_', ' ').title()}:")
                            for entry in detail:
                                if isinstance(entry, dict):
                                    nested = _flatten_value(entry)
                                    for line in nested.split("\n"):
                                        lines.append(f"- {line}")
                                else:
                                    lines.append(f"- {entry}")
                        elif isinstance(detail, dict):
                            lines.append(f"{key.replace('_', ' ').title()}:" )
                            nested = _flatten_value(detail)
                            for line in nested.split("\n"):
                                lines.append(f"- {line}")
                        else:
                            lines.append(f"{key.replace('_', ' ').title()}: {detail}")
                else:
                    lines.append(f"- {item}")
        lines.append("")

    return "\n".join(lines)


def _maybe_format_bold_lead(line: str, para):
    if ":" in line and not line.startswith("- "):
        key, rest = line.split(":", 1)
        run = para.add_run(f"{key.strip()}: ")
        run.bold = True
        para.add_run(rest.strip())
        return True
    return False


def _is_date_line(line: str) -> bool:
    return any(token in line.lower() for token in ["jan", "feb", "mar", "apr", "may", "jun", "jul", "aug", "sep", "oct", "nov", "dec", "20", "202"])


def create_docx(resume_text: str) -> BytesIO:
    """Return a BytesIO buffer containing a polished formatted .docx resume."""
    structured = _load_structured_resume(resume_text)
    if structured:
        resume_text = _resume_dict_to_text(structured)

    resume_text = _clean_text(resume_text)
    doc = Document()
    default_style = doc.styles["Normal"]
    default_style.font.name = "Calibri"
    default_style.font.size = Pt(11)

    header_lines, sections = _split_header_and_sections(resume_text)
    header = _parse_resume_header(header_lines)

    if header["name"]:
        paragraph = doc.add_paragraph()
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run = paragraph.add_run(header["name"])
        run.bold = True
        run.font.size = Pt(28)

    if header["title"]:
        paragraph = doc.add_paragraph()
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run = paragraph.add_run(header["title"])
        run.italic = True
        run.font.size = Pt(12)

    if header["contact"]:
        paragraph = doc.add_paragraph()
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run = paragraph.add_run(" | ".join(header["contact"]))
        run.font.size = Pt(10)

    if header["extras"]:
        extras_paragraph = doc.add_paragraph()
        extras_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run = extras_paragraph.add_run(" | ".join(header["extras"]))
        run.font.size = Pt(10)

    if header["name"] or header["title"] or header["contact"] or header["extras"]:
        doc.add_paragraph()

    for section in sections:
        title = section["title"]
        if title:
            paragraph = doc.add_paragraph()
            run = paragraph.add_run(title.upper())
            run.bold = True
            paragraph.style = doc.styles["Heading 2"]

        for line in section["items"]:
            if line.startswith("- ") or line.startswith("• "):
                doc.add_paragraph(line[2:].strip(), style="List Bullet")
            elif _is_date_line(line) and "|" not in line and any(part.isdigit() for part in line.split()):
                paragraph = doc.add_paragraph()
                paragraph.add_run(line)
            elif ":" in line and not line.startswith(("- ", "• ")):
                paragraph = doc.add_paragraph()
                if not _maybe_format_bold_lead(line, paragraph):
                    paragraph.add_run(line)
            elif "|" in line and len(line) <= 120:
                parts = [part.strip() for part in line.split("|") if part.strip()]
                para = doc.add_paragraph()
                first = True
                for part in parts:
                    if not first:
                        para.add_run("   ")
                    run = para.add_run(part)
                    run.bold = True
                    first = False
            else:
                doc.add_paragraph(line)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

def create_pdf(resume_text: str) -> BytesIO:
    """
    Return a BytesIO buffer containing a highly ATS-optimized PDF resume.

    ATS principles applied:
    - Single-column, top-to-bottom linear layout (no tables, no text boxes)
    - Standard fonts only (Helvetica) — no embedded/decorative fonts
    - No headers/footers, no images, no graphics, no watermarks
    - Logical section order with plain-text headings in ALL CAPS
    - Bullet points rendered as plain hyphens (most ATS-safe)
    - No hyperlink formatting that breaks text extraction
    - Metadata set for proper document identification
    - Generous margins for clean text extraction
    """

    # ── 1. Pre-process input ──────────────────────────────────────────────
    structured = _load_structured_resume(resume_text)
    if structured:
        resume_text = _resume_dict_to_text(structured)

    resume_text = _clean_text(resume_text)
    header_lines, sections = _split_header_and_sections(resume_text)
    header = _parse_resume_header(header_lines)

    # ── 2. Document setup ─────────────────────────────────────────────────
    buffer = BytesIO()
    PAGE_W, PAGE_H = letter
    MARGIN = 0.75 * inch          # 0.75 in — ATS parsers need breathing room

    doc = SimpleDocTemplate(
        buffer,
        pagesize=letter,
        leftMargin=MARGIN,
        rightMargin=MARGIN,
        topMargin=0.6 * inch,
        bottomMargin=0.6 * inch,
        # Set PDF metadata so ATS can identify the document
        title=header.get("name", "Resume"),
        author=header.get("name", ""),
        subject="Resume",
        creator="Resume Builder",
    )

    # ── 3. Style definitions ──────────────────────────────────────────────
    # ATS rule: use only standard font families (Helvetica/Times/Courier)
    FONT_REGULAR  = "Helvetica"
    FONT_BOLD     = "Helvetica-Bold"
    BLACK         = colors.HexColor("#000000")
    RULE_COLOR    = colors.HexColor("#999999")

    def style(name, **kw) -> ParagraphStyle:
        base = dict(
            fontName=FONT_REGULAR,
            fontSize=11,
            leading=15,
            textColor=BLACK,
            alignment=TA_LEFT,
            spaceAfter=0,
            spaceBefore=0,
        )
        base.update(kw)
        return ParagraphStyle(name, **base)

    st_name = style(
        "Name",
        fontName=FONT_BOLD,
        fontSize=20,
        leading=24,
        alignment=TA_CENTER,
        spaceAfter=2,
    )
    st_title = style(
        "Title",
        fontSize=11,
        leading=14,
        alignment=TA_CENTER,
        spaceAfter=2,
    )
    st_contact = style(
        "Contact",
        fontSize=10,
        leading=13,
        alignment=TA_CENTER,
        spaceAfter=6,
        textColor=colors.HexColor("#333333"),
    )
    st_section = style(
        "Section",
        fontName=FONT_BOLD,
        fontSize=11,
        leading=14,
        spaceBefore=10,
        spaceAfter=2,
        # No background / shading — ATS strips visuals
    )
    st_job_header = style(
        "JobHeader",
        fontName=FONT_BOLD,
        fontSize=11,
        leading=14,
        spaceBefore=4,
        spaceAfter=1,
    )
    st_date = style(
        "Date",
        fontSize=10,
        leading=13,
        textColor=colors.HexColor("#444444"),
        spaceAfter=1,
    )
    st_body = style(
        "Body",
        fontSize=11,
        leading=15,
        spaceAfter=1,
    )
    st_bullet = style(
        # ATS rule: plain hyphens, not Unicode bullets (•/▪/→)
        "Bullet",
        fontSize=11,
        leading=15,
        leftIndent=14,
        firstLineIndent=-14,   # hanging indent
        spaceAfter=1,
    )

    # ── 4. Build story ────────────────────────────────────────────────────
    story = []

    def hr(thickness=0.5, color=RULE_COLOR, space_before=2, space_after=4):
        story.append(Spacer(1, space_before))
        story.append(HRFlowable(
            width="100%",
            thickness=thickness,
            color=color,
            spaceAfter=space_after,
        ))

    def safe_para(text: str, st: ParagraphStyle) -> Paragraph:
        """Escape XML special chars so ReportLab doesn't choke."""
        text = (text
                .replace("&", "&amp;")
                .replace("<", "&lt;")
                .replace(">", "&gt;"))
        return Paragraph(text, st)

    # ── 4a. Header block ──────────────────────────────────────────────────
    header_block = []

    name = header.get("name", "").strip()
    if name:
        header_block.append(safe_para(name, st_name))

    title = header.get("title", "").strip()
    if title:
        header_block.append(safe_para(title, st_title))

    contacts = header.get("contact", [])
    if contacts:
        # ATS rule: plain text separated by " | ", no clickable links
        contact_line = "  |  ".join(contacts)
        header_block.append(safe_para(contact_line, st_contact))

    if header_block:
        story.append(KeepTogether(header_block))
        hr(thickness=1.0, color=BLACK, space_before=4, space_after=6)

    # ── 4b. Sections ──────────────────────────────────────────────────────
    for section in sections:
        sec_title = (section.get("title") or "").strip()
        items     = section.get("items", [])

        if not sec_title and not items:
            continue

        block = []

        if sec_title:
            # ATS rule: ALL-CAPS section headings improve parser recognition
            block.append(safe_para(sec_title.upper(), st_section))
            block.append(HRFlowable(
                width="100%",
                thickness=0.4,
                color=RULE_COLOR,
                spaceAfter=3,
            ))

        for line in items:
            line = line.strip()
            if not line:
                continue

            # Bullet line
            if line.startswith(("- ", "* ", "• ", "· ")):
                # ATS rule: use plain hyphen, not unicode bullet
                clean = line[2:].strip()
                block.append(safe_para(f"- {clean}", st_bullet))

            # Date / tenure line  (e.g. "Jan 2020 – Dec 2022")
            elif _is_date_line(line) and any(ch.isdigit() for ch in line):
                block.append(safe_para(line, st_date))

            # Job/role header line  (e.g. "Software Engineer | Acme Corp")
            elif "|" in line and len(line) <= 120:
                # ATS rule: replace fancy pipes with plain text separator
                clean = line.replace("|", " - ")
                block.append(safe_para(clean, st_job_header))

            # Plain prose / everything else
            else:
                block.append(safe_para(line, st_body))

        block.append(Spacer(1, 6))
        story.append(KeepTogether(block))

    # ── 5. Build PDF ──────────────────────────────────────────────────────
    doc.build(story)
    buffer.seek(0)
    return buffer
# def create_pdf(resume_text: str) -> BytesIO:
    """Return a BytesIO buffer containing a polished formatted PDF resume."""
    structured = _load_structured_resume(resume_text)
    if structured:
        resume_text = _resume_dict_to_text(structured)

    resume_text = _clean_text(resume_text)
    header_lines, sections = _split_header_and_sections(resume_text)
    header = _parse_resume_header(header_lines)

    buffer = BytesIO()
    width, height = letter
    c = canvas.Canvas(buffer, pagesize=letter)
    LEFT_MARGIN = 50
    RIGHT_MARGIN = 50
    usable_width = width - LEFT_MARGIN - RIGHT_MARGIN
    y = height - 80

    # Top header area
    c.setFillColor(colors.HexColor("#ffffff"))
    c.rect(0, height - 140, width, 140, fill=1, stroke=0)

    if header["name"]:
        c.setFillColor(colors.HexColor("#111111"))
        c.setFont("Helvetica-Bold", 26)
        c.drawString(LEFT_MARGIN, height - 60, header["name"])
        y = height - 90
    else:
        c.setFillColor(colors.HexColor("#111111"))
        c.setFont("Helvetica-Bold", 26)
        c.drawString(LEFT_MARGIN, height - 60, "ATS Optimized Resume")
        y = height - 90

    if header["title"]:
        c.setFont("Helvetica", 12)
        c.drawString(LEFT_MARGIN, y, header["title"])
        y -= 18

    if header["contact"]:
        c.setFont("Helvetica", 9)
        contact_line = " | ".join(header["contact"])
        c.drawString(LEFT_MARGIN, y, contact_line)
        y -= 20

    c.setStrokeColor(colors.HexColor("#111111"))
    c.setLineWidth(1)
    c.line(LEFT_MARGIN, y, width - RIGHT_MARGIN, y)
    y -= 24

    def draw_wrapped(text: str, x: float, y_pos: float, font_name: str, font_size: int, leading: int = 16):
        c.setFont(font_name, font_size)
        words = text.split()
        line = ""
        for word in words:
            test = f"{line}{word} "
            if c.stringWidth(test, font_name, font_size) <= usable_width:
                line = test
            else:
                c.drawString(x, y_pos, line.strip())
                y_pos -= leading
                line = f"{word} "
                if y_pos < 70:
                    c.showPage()
                    y_pos = height - 80
                    c.setFont(font_name, font_size)
        if line:
            c.drawString(x, y_pos, line.strip())
            y_pos -= leading
        return y_pos

    for section in sections:
        title = section["title"]
        if title:
            c.setFillColor(colors.HexColor("#111111"))
            c.setFont("Helvetica-Bold", 14)
            y -= 8
            c.drawString(LEFT_MARGIN, y, title.upper())
            y -= 16
            c.setLineWidth(0.75)
            c.setStrokeColor(colors.HexColor("#d1d5db"))
            c.line(LEFT_MARGIN, y + 8, width - RIGHT_MARGIN, y + 8)
            y -= 12

        for line in section["items"]:
            if line.startswith("- ") or line.startswith("• "):
                c.setFillColor(colors.HexColor("#111111"))
                y = draw_wrapped(f"• {line[2:].strip()}", LEFT_MARGIN + 10, y, "Helvetica", 11, leading=16)
            elif _is_date_line(line) and any(char.isdigit() for char in line):
                c.setFillColor(colors.HexColor("#111111"))
                c.setFont("Helvetica-Bold", 11)
                c.drawString(LEFT_MARGIN, y, line)
                y -= 16
            elif "|" in line and len(line) <= 120:
                c.setFillColor(colors.HexColor("#111111"))
                y = draw_wrapped(line.replace("|", "   "), LEFT_MARGIN, y, "Helvetica-Bold", 11, leading=16)
                y -= 6
            else:
                c.setFillColor(colors.HexColor("#111111"))
                y = draw_wrapped(line, LEFT_MARGIN, y, "Helvetica", 11, leading=16)
            if y < 70:
                c.showPage()
                y = height - 80

        y -= 16
        if y < 70:
            c.showPage()
            y = height - 80

    c.save()
    buffer.seek(0)
    return buffer
